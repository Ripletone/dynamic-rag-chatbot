import streamlit as st

# --- Page Configuration (MUST BE THE FIRST STREAMLIT COMMAND) ---
st.set_page_config(page_title="Dynamic RAG Chatbot with Memory", layout="wide")

import os
from operator import itemgetter
import shutil
import re
from urllib.parse import urlparse

# CRITICAL FIX: Use nest_asyncio to handle event loop conflicts in Streamlit
import nest_asyncio
nest_asyncio.apply()

# LangChain Imports - NOW USING FAISS INSTEAD OF CHROMA
from langchain_huggingface.embeddings import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.documents import Document
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough, RunnableParallel
from langchain_core.messages import HumanMessage, AIMessage
from langchain_core.tools import Tool # Re-importing Tool from langchain_core for explicit usage

# DIRECT IMPORTS for Document Loading and Tools (replaces langchain_community)
import pypdf
import requests
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
import markdown # if using markdown files
from duckduckgo_search import DDGS

# FAISS direct import
import faiss
import numpy as np


# --- Configuration Constants ---
GEMINI_MODEL_NAME = "gemini-1.5-flash-latest"
EMBEDDING_MODEL_NAME = "sentence-transformers/all-MiniLM-L6-v2"
DEFAULT_COLLECTION_NAME = "default_faiss_collection"


# --- API Key Handling ---
try:
    GOOGLE_API_KEY = st.secrets["GOOGLE_API_KEY"]
except KeyError:
    st.error("`GOOGLE_API_KEY` not found in `.streamlit/secrets.toml`. Please add it to your secrets file.")
    st.stop()


# --- Streamlit Session State Initialization ---
def initialize_session_state():
    if "vector_db" not in st.session_state:
        st.session_state.vector_db = None
    if "current_content_source" not in st.session_state:
        st.session_state.current_content_source = None
    if "current_collection_name" not in st.session_state:
        st.session_state.current_collection_name = DEFAULT_COLLECTION_NAME
    if "messages" not in st.session_state:
        st.session_state.messages = []
    # Re-introducing url_input_key as a counter for text_input's key
    if "url_input_key" not in st.session_state:
        st.session_state.url_input_key = 0
    if "uploaded_file_key" not in st.session_state:
        st.session_state.uploaded_file_key = 0

    if "llm_temperature" not in st.session_state:
        st.session_state.llm_temperature = 0.0
    if "top_k_retrieval" not in st.session_state:
        st.session_state.top_k_retrieval = 4
    if "selected_search_type" not in st.session_state:
        st.session_state.selected_search_type = "Similarity"
    if "fetch_k_mmr" not in st.session_state:
        st.session_state.fetch_k_mmr = 20
    if "lambda_mult_mmr" not in st.session_state:
        st.session_state.lambda_mult_mmr = 0.5

    if "faiss_indexes" not in st.session_state:
        st.session_state.faiss_indexes = {}

    if "last_loaded_url" not in st.session_state:
        st.session_state["last_loaded_url"] = ""
    # New session state variable to reliably store the URL submitted by the button
    if "submitted_url" not in st.session_state:
        st.session_state.submitted_url = ""


initialize_session_state()


# --- Helper Functions for Naming and Data Processing ---

def clean_collection_name(name: str) -> str:
    """
    Cleans a string to be a valid collection identifier for FAISS.
    FAISS doesn't have strict naming like ChromaDB, but a cleaned name is good for keys.
    """
    if not isinstance(name, str) or not name.strip():
        return DEFAULT_COLLECTION_NAME

    cleaned_name = re.sub(r'[^a-zA-Z0-9._-]', '_', name)
    cleaned_name = re.sub(r'^[^a-zA-Z0-9]+', '', cleaned_name)
    cleaned_name = re.sub(r'[^a-zA-Z0-9]+$', '', cleaned_name)
    if len(cleaned_name) < 3:
        cleaned_name = "collection_" + cleaned_name.lower()
    cleaned_name = cleaned_name[:100]

    if not cleaned_name.strip() or not cleaned_name[0].isalnum() or not cleaned_name[-1].isalnum():
        return DEFAULT_COLLECTION_NAME

    return cleaned_name.lower()

def get_url_collection_name(url: str) -> str:
    """Generates a consistent collection name from a URL."""
    parsed_url = urlparse(url)
    base_name = f"{parsed_url.netloc}{parsed_url.path}".strip('/')
    if not base_name:
        base_name = "web_content"
    unique_suffix = str(abs(hash(url)))
    combined_name = f"{base_name}_{unique_suffix}"
    return clean_collection_name(combined_name)


# --- Custom Document Loading Functions (Replacing langchain_community.document.loaders) ---

def load_documents_from_file(file_path: str) -> list[Document]:
    """Loads content from a file using direct library calls."""
    documents = []
    content = ""
    try:
        file_extension = os.path.splitext(file_path)[1].lower()
        if file_extension == ".pdf":
            with open(file_path, "rb") as f:
                reader = pypdf.PdfReader(f)
                for page in reader.pages:
                    content += page.extract_text() or ""
        elif file_extension == ".md":
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
        elif file_extension == ".docx":
            doc = DocxDocument(file_path)
            for para in doc.paragraphs:
                content += para.text + "\n"
        elif file_extension == ".txt":
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
        else:
            st.error(f"Unsupported file type: {file_extension}. Only PDF, TXT, MD, DOCX supported for direct loading.")
            return []
        documents.append(Document(page_content=content, metadata={"source": file_path, "file_name": os.path.basename(file_path)}))
    except Exception as e:
        st.error(f"Error loading file {file_path}: {e}")
        return []
    return documents

def load_documents_from_url(url: str) -> list[Document]:
    """Loads content from a URL using direct library calls."""
    documents = []
    try:
        response = requests.get(url, timeout=10)
        response.raise_for_status() # Raise an exception for HTTP errors
        soup = BeautifulSoup(response.text, 'html.parser')
        # Extract readable text from the HTML, removing script/style tags
        for script_or_style in soup(["script", "style"]):
            script_or_style.extract()
        text_content = soup.get_text(separator=' ', strip=True)
        documents.append(Document(page_content=text_content, metadata={"source": url}))
    except requests.exceptions.RequestException as e:
        st.error(f"Error fetching URL {url}: {e}")
        return []
    except Exception as e:
        st.error(f"Error parsing URL content from {url}: {e}")
        return []
    return documents


# --- Caching and Document Handling Functions ---

@st.cache_resource
def get_embedding_function():
    """Loads and caches the embedding model (PURE function, NO UI interactions)."""
    embeddings = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL_NAME, model_kwargs={'device': 'cpu'})
    return embeddings


# This CustomFAISSVectorStore class is a simplified wrapper for direct faiss-cpu usage,
# to mimic some behavior of langchain-community's FAISS, specifically for retrieval.
# It's NOT a full replacement for all langchain-community.vectorstores.FAISS methods
# like 'from_documents' or 'as_retriever' directly.
# The 'as_retriever' method here returns a custom RetrieverWrapper.
class CustomFAISSVectorStore:
    def __init__(self, embeddings_model, documents=None, dim=None):
        self.embeddings_model = embeddings_model
        self.index = None
        self.doc_store = [] # To store original documents and their content for retrieval
        self.dim = dim # Dimension of embeddings

        if documents:
            self.add_documents(documents)

    def add_documents(self, documents: list[Document]):
        if not documents:
            return

        # Get embeddings for new documents
        texts = [doc.page_content for doc in documents]
        # Ensure embeddings are numpy arrays
        new_embeddings = np.array(self.embeddings_model.embed_documents(texts), dtype=np.float32)

        if self.index is None:
            self.dim = new_embeddings.shape[1]
            self.index = faiss.IndexFlatL2(self.dim) # L2 distance (Euclidean)
            self.doc_store = documents # Initialize doc_store with the first set of documents
        else:
            self.doc_store.extend(documents) # Extend if index already exists

        # Add vectors to the FAISS index
        self.index.add(new_embeddings)

    def similarity_search_with_score(self, query: str, k: int = 4) -> list[tuple[Document, float]]:
        query_embedding = np.array(self.embeddings_model.embed_query(query), dtype=np.float32).reshape(1, -1)
        
        if self.index is None or self.index.ntotal == 0:
            return []
        
        distances, indices = self.index.search(query_embedding, k)
        
        results = []
        for i, dist in zip(indices[0], distances[0]):
            if i >= 0 and i < len(self.doc_store): # Ensure index is valid and within bounds
                results.append((self.doc_store[i], float(dist)))
        return results

    def as_retriever(self, search_type="similarity", search_kwargs={}):
        # This is a simplified retriever to avoid direct langchain.vectorstores.FAISS dependency.
        # It won't support all features of LangChain's native FAISS retriever.
        # For agent functionality, this part will need the LangChain ecosystem.
        if search_type == "similarity":
            k = search_kwargs.get("k", 4)
            return RetrieverWrapper(self, k)
        else:
            st.warning("Only 'similarity' search type is supported by the custom FAISS wrapper. Defaulting to similarity.")
            return RetrieverWrapper(self, search_kwargs.get("k", 4))

class RetrieverWrapper:
    # A simple wrapper to make CustomFAISSVectorStore usable as a retriever
    def __init__(self, vector_store_instance, k):
        self.vector_store_instance = vector_store_instance
        self.k = k

    def get_relevant_documents(self, query: str) -> list[Document]:
        results_with_score = self.vector_store_instance.similarity_search_with_score(query, k=self.k)
        # Return only the documents (discard scores for this interface)
        return [doc for doc, score in results_with_score]


def _load_or_create_faiss_index(text_chunks, collection_name_for_cache):
    """
    Internal function to create or load a FAISS vector store.
    This function directly manages st.session_state.faiss_indexes.
    """
    embedding_function = get_embedding_function()

    if text_chunks:
        try:
            # Using the custom FAISS wrapper here
            faiss_index = CustomFAISSVectorStore(
                embeddings_model=embedding_function,
                documents=text_chunks
            )
            st.session_state.faiss_indexes[collection_name_for_cache] = faiss_index
            st.toast(f"Knowledge base collection '{collection_name_for_cache}' created/updated successfully with {len(text_chunks)} chunks!", icon="✨")
            st.session_state.current_content_source = (
                f"Updated from '{collection_name_for_cache}'" if collection_name_for_cache != DEFAULT_COLLECTION_NAME else "Default Collection"
            )
            return faiss_index
        except Exception as e:
            st.error(f"Error creating/updating knowledge base with FAISS: {e}")
            return None
    else:
        faiss_index = st.session_state.faiss_indexes.get(collection_name_for_cache)

        if faiss_index:
            st.toast(f"Knowledge base collection '{collection_name_for_cache}' loaded from session state!", icon="📚")
            return faiss_index
        else:
            st.warning(f"No existing knowledge base collection '{collection_name_for_cache}' in memory. Please upload a document or load a URL.")
            return None


def manage_vector_store(text_chunks=None, collection_name=DEFAULT_COLLECTION_NAME):
    """
    Manages the active FAISS vector store in session state.
    Called on initial load, file upload, or URL load.
    """
    cleaned_collection_name = clean_collection_name(collection_name)
    st.session_state.current_collection_name = cleaned_collection_name

    with st.spinner("Loading embedding model (if not already cached)..."):
        _ = get_embedding_function()
    st.toast("Embedding model ready!", icon="✅")

    if text_chunks is not None and len(text_chunks) > 0:
        st.info(f"Processing content for in-memory knowledge base '{cleaned_collection_name}'...")
        vector_db = _load_or_create_faiss_index(text_chunks, cleaned_collection_name)
        st.session_state.vector_db = vector_db
        if vector_db:
             st.session_state.current_content_source = (
                f"Loaded from '{collection_name}'"
            )
        else:
            st.session_state.current_content_source = None
    else:
        vector_db = st.session_state.faiss_indexes.get(cleaned_collection_name) # Try to load from existing in-memory indexes
        st.session_state.vector_db = vector_db
        if vector_db:
            st.session_state.current_content_source = (
                f"Pre-existing knowledge base (in-memory: {cleaned_collection_name})"
            )
        else:
            st.session_state.current_content_source = None


@st.cache_data(show_spinner=False)
def process_document_to_chunks(uploaded_file, chunk_size, chunk_overlap):
    """Loads and splits an uploaded document into text chunks using direct loaders."""
    if uploaded_file is None:
        return []

    temp_file_dir = "./temp_uploaded_files"
    os.makedirs(temp_file_dir, exist_ok=True)
    temp_file_path = os.path.join(temp_file_dir, uploaded_file.name)

    try:
        with open(temp_file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())

        st.info(f"Loading document: {uploaded_file.name}")

        # Use custom direct loading function
        documents = load_documents_from_file(temp_file_path)

        if not documents:
            st.warning(f"No content found in {uploaded_file.name}.")
            return []

        text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        text_chunks = text_splitter.split_documents(documents)

        st.toast(f"Document '{uploaded_file.name}' processed into {len(text_chunks)} chunks!", icon="📄")
        return text_chunks

    except Exception as e:
        st.error(f"Error processing document '{uploaded_file.name}': {e}")
        return []
    finally:
        if os.path.exists(temp_file_path):
            os.remove(temp_file_path)


@st.cache_data(show_spinner=False)
def process_url_to_chunks(url_input, chunk_size, chunk_overlap):
    """Loads and splits content from a URL into text chunks using direct loaders."""
    if not url_input:
        return []

    with st.status(f"Loading content from URL: {url_input}...", expanded=True) as status_message:
        try:
            status_message.write("1. Fetching content from URL...")
            # Use custom direct loading function
            documents = load_documents_from_url(url_input)

            if not documents:
                raise ValueError("No content extracted from URL.")
            status_message.write(f"2. Extracted {len(documents)} document(s).")

            status_message.write("3. Splitting text into chunks...")
            text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
            text_chunks = text_splitter.split_documents(documents)
            if not text_chunks:
                raise ValueError("No text chunks generated from content.")
            status_message.write(f"4. Created {len(text_chunks)} text chunks.")
            status_message.update(label=f"Content from '{url_input}' processed!", state="complete", expanded=False)
            return text_chunks
        except Exception as e:
            status_message.update(label=f"Error loading URL: {e}", state="error", expanded=True)
            st.error(f"Error loading URL: {e}")
            return []

# --- IMPORTANT: Agent and Chain Dependencies ---
# The following functions (get_llm_agent, generate_answer_with_memory using AgentExecutor,
# create_tool_calling_agent, create_retriever_tool) are part of the main 'langchain'
# package which conflicts with 'langchain-google-genai 2.x' due to 'langchain-core' versions.
#
# To make this code run with 'langchain-google-genai==2.1.5' (for Gemini 1.5-flash-latest),
# you MUST ensure 'langchain' (the main package) is NOT in your requirements.txt.
#
# If 'langchain' is removed, these functions will cause ModuleNotFoundError.
# Replacing them requires significant refactoring to manually implement agent logic,
# tool calling, and chain orchestration using only 'langchain_core' or direct API calls.
#
# For now, I'm leaving this section with a fallback mechanism.
# The 'try-except ImportError' blocks will attempt to use LangChain's agent components.
# If they fail (because 'langchain' is not installed), a simpler RAG chain will be used,
# which will NOT use tools dynamically (only the knowledge base if available).

# Agent for flexible tool use
def get_llm_agent(
    vector_db,
    llm_temperature_param: float,
    search_type_param: str,
    k_param: int,
    fetch_k_param: int,
    lambda_mult_param: float
):
    """
    Constructs and returns an LLM Agent capable of using a RAG retriever and a web search tool.
    NOTE: This function attempts to use LangChain Agent components which might conflict with
    langchain-google-genai 2.x unless the 'langchain' package is compatible or removed.
    If 'langchain' is not installed, it falls back to a simpler RAG chain without dynamic tool use.
    """
    llm = ChatGoogleGenerativeAI(
        model=GEMINI_MODEL_NAME,
        temperature=llm_temperature_param,
        google_api_key=GOOGLE_API_KEY
    )

    tools = []

    # --- Knowledge Base Tool ---
    if vector_db:
        try:
            # Attempt to import from the full langchain package
            from langchain.tools.retriever import create_retriever_tool
            retriever = vector_db.as_retriever(search_type="similarity", search_kwargs={"k": k_param})

            if search_type_param == "MMR":
                st.warning("MMR search type is not directly supported by FAISS retriever. Using 'similarity' instead.")

            knowledge_base_tool = create_retriever_tool(
                retriever,
                name="KnowledgeBase_Search",
                description="Searches and returns information from the user-provided documents in the knowledge base. Use this for questions specifically about uploaded content.",
            )
            tools.append(knowledge_base_tool)
        except ImportError:
            # Fallback: create a manual tool using the custom retriever
            st.info("LangChain's 'create_retriever_tool' not found. Using a basic custom retriever tool.")
            knowledge_base_tool = Tool(
                name="KnowledgeBase_Search",
                func=lambda query: "\n".join([doc.page_content for doc in vector_db.as_retriever(search_kwargs={"k":k_param}).get_relevant_documents(query)]),
                description="Searches and returns information from the user-provided documents in the knowledge base. Use this for questions specifically about uploaded content."
            )
            tools.append(knowledge_base_tool)


    # --- Direct DuckDuckGo Search Tool (replaces langchain_community.tools.DuckDuckGoSearchResults) ---
    def duckduckgo_search_tool_func(query: str) -> str:
        with DDGS() as ddgs:
            results = [r for r in ddgs.text(query, max_results=5)]
            if results:
                return "\n".join([f"Title: {r['title']}\nURL: {r['href']}\nSnippet: {r['snippet']}" for r in results])
            return "No search results found."

    web_search_tool = Tool(
        name="Web_Search",
        description="Useful for when you need to answer questions about current events, facts, or anything not covered in the provided documents. Prioritize the KnowledgeBase_Search tool if the question is likely about uploaded content.",
        func=duckduckgo_search_tool_func,
    )
    tools.append(web_search_tool)

    if not tools:
        st.warning("No tools are available for the agent. Please load a document or ensure tools are correctly defined.")
        return None

    agent_prompt = ChatPromptTemplate.from_messages([
        ("system", """You are a helpful AI assistant. Your primary goal is to answer user questions accurately and comprehensively.

        Always prioritize 'KnowledgeBase_Search' if the question can be answered from the documents.
        If you use 'Web_Search', aim to summarize relevant information and indicate that it came from the web (e.g., "According to a web search...").
        If you use 'KnowledgeBase_Search', indicate that it came from the knowledge base (e.g., "From the knowledge base...").
        If you cannot find an answer using any tool, state that you don't know.
        """),

        MessagesPlaceholder(variable_name="chat_history"),
        ("human", "{input}"),
        MessagesPlaceholder(variable_name="agent_scratchpad")
    ])

    # --- Agent Execution Logic ---
    # This part requires 'langchain' (the main package). If removed, this will fail.
    # You would need to implement a custom agent loop or switch to a different agent framework.
    try:
        from langchain.agents import AgentExecutor, create_tool_calling_agent
        agent = create_tool_calling_agent(llm, tools, agent_prompt)
        agent_executor = AgentExecutor(
            agent=agent,
            tools=tools,
            verbose=True,
            handle_parsing_errors=True
        )
        return agent_executor
    except ImportError:
        st.error("LangChain agent components (like AgentExecutor, create_tool_calling_agent) are missing or incompatible.")
        st.info("Falling back to a basic RAG chain. Dynamic tool use (like Web Search) will be disabled.")
        # Fallback for simple RAG chain if agent is not available.
        # This will NOT use tools dynamically but only do RAG if vector_db exists.
        if vector_db:
            retriever = vector_db.as_retriever(search_kwargs={"k": k_param})

            # FIX: Restructure the RAG chain to correctly use the RetrieverWrapper
            # The 'context' key will now get documents from the retriever based on 'input'
            rag_chain_input = RunnableParallel(
                context=itemgetter("input") | retriever.get_relevant_documents,
                question=itemgetter("input"),
                chat_history=itemgetter("chat_history")
            )

            rag_chain = (
                rag_chain_input
                | ChatPromptTemplate.from_messages([
                    ("system", "You are a helpful AI assistant. Use the following retrieved context to answer the question. If the context is not sufficient, state that you don't know.\n\nContext: {context}"),
                    MessagesPlaceholder(variable_name="chat_history"),
                    ("human", "{question}")
                ])
                | llm
                | StrOutputParser()
            )
            return rag_chain
        else:
            return None # No RAG or agent available without necessary components


def generate_answer_with_memory(
    query,
    chat_history,
    vector_db
):
    """Generates an answer using the LLM agent/chain, including conversational memory and tool use."""

    llm_temperature_param = st.session_state.llm_temperature
    search_type_param = st.session_state.selected_search_type
    k_param = st.session_state.top_k_retrieval
    fetch_k_param = st.session_state.fetch_k_mmr
    lambda_mult_param = st.session_state.lambda_mult_mmr

    agent_or_chain = get_llm_agent( # This might return an AgentExecutor or a Runnable chain
        vector_db,
        llm_temperature_param,
        search_type_param,
        k_param,
        fetch_k_param,
        lambda_mult_param
    )

    if agent_or_chain is None:
        st.warning("Agent or basic RAG chain could not be initialized. Please ensure necessary components are available.")
        return "An error occurred, no LLM functionality initialized.", []

    processed_query = str(query).strip()
    if not processed_query:
        return "I need a question to answer!", []

    formatted_chat_history = []
    for msg in chat_history:
        if msg["role"] == "user":
            formatted_chat_history.append(HumanMessage(content=str(msg["content"])))
        elif msg["role"] == "assistant":
            formatted_chat_history.append(AIMessage(content=str(msg["content"])))

    try:
        # Check if it's an AgentExecutor or a simple Runnable (chain)
        if hasattr(agent_or_chain, 'invoke'):
            response = agent_or_chain.invoke({
                "input": processed_query,
                "chat_history": formatted_chat_history
            })
        else:
            # If it's not a standard LangChain agent/chain, it might be a simpler callable.
            # Given the fallback is a Runnable, this branch should be fine.
            response = agent_or_chain.invoke({"input": processed_query, "chat_history": formatted_chat_history})


        # The structure of response depends on whether AgentExecutor or a simple Runnable was returned.
        response_text = ""
        if isinstance(response, dict) and "output" in response:
            response_text = response["output"]
        elif isinstance(response, str):
            response_text = response
        else:
            response_text = f"Unexpected response format: {type(response)} - {response}"


        source_docs_for_display = [] # Agent does not directly return source docs in this setup, but you could parse verbose output
        return response_text, source_docs_for_display

    except Exception as e:
        st.error(f"Error generating answer with agent/chain: {e}")
        return "An error occurred while trying to generate an answer. Please try again.", []


# --- Main Streamlit App Layout and Logic ---

st.markdown(
    """
    <style>
    /* Force body and general text color/font */
    body, p, div, span, h1, h2, h3, h4, h5, h6, .stMarkdown {
        color: #000000 !important;
        font-family: sans-serif !important;
        font-weight: normal !important;
    }

    /* Keep the original styling for Streamlit-generated code blocks and preformatted text */
    code, pre {
        font-family: monospace !important;
    }

    /* Styling for the URL input text box */
    div[data-testid="stTextInput"] input {
        border: 1px solid #4CAF50;
        border-radius: 8px;
        padding: 10px 15px;
        background-color: #f8fcf8;
        color: #333 !important;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        font-family: sans-serif !important;
        font-weight: normal !important;
    }

    /* Styling for the chat input text box */
    div[data-testid="stChatInput"] input {
        border: 1px solid #007bff;
        border-radius: 20px;
        padding: 12px 20px;
        background-color: #eaf6ff;
        color: #333 !important;
        box-shadow: 0 2px 5px rgba(0,0,0,0.15);
        font-family: sans-serif !important;
        font-weight: normal !important;
    }

    /* Styling for the Sources expander */
    .stExpander {
        border: 1px solid #a8dadc;
        border-radius: 10px;
        padding: 5px 15px;
        margin-top: 20px;
        background-color: #f1f8f9;
        box-shadow: 0 1px 3px rgba(0,0,0,0.08);
    }

    /* Optional: Style for the chat message boxes themselves */
    .stChatMessage {
        background-color: #ffffff;
        border-radius: 12px;
        padding: 15px;
        margin-bottom: 10px;
        box-shadow: 0 1px 3px rgba(0,0,0,0.05);
    }

    .stChatMessage.st-chat-message-user {
        background-color: #e0f2f7;
        text-align: right;
    }
    .stChatMessage.st-chat-message-assistant {
        background-color: #e6ffe6;
        text-align: left;
    }

    </style>
    """,
    unsafe_allow_html=True
)

st.title("Dynamic RAG Chatbot with Memory")
st.write("Upload a document or load a URL to create a knowledge base, then ask questions about it, or ask general questions!")


# Sidebar for controls
st.sidebar.header("Configuration")

uploaded_file = st.sidebar.file_uploader("Upload a Document", type=["pdf", "txt", "md", "docx"], key="file_uploader")

# --- Document Processing Settings (using expander and columns) ---
with st.sidebar.expander("Document Processing Settings", expanded=True):
    st.markdown("Adjust how documents are split into chunks for retrieval.")
    col1, col2 = st.columns(2)

    with col1:
        st.session_state.chunk_size = st.slider(
            "Chunk Size",
            min_value=100, max_value=2000, value=1000, step=50,
            help="Size of text segments (characters)."
        )
    with col2:
        st.session_state.chunk_overlap = st.slider(
            "Overlap",
            min_value=0, max_value=500, value=200, step=25,
            help="Overlapping characters between chunks."
        )

# --- Language Model Settings (using expander) ---
with st.sidebar.expander("Language Model Settings", expanded=True):
    st.markdown("Configure the behavior of the AI model.")
    st.session_state.llm_temperature = st.slider(
        "LLM Temperature",
        min_value=0.0,
        max_value=1.0,
        value=0.0,
        step=0.05,
        help="Controls creativity. Lower = more deterministic; higher = more creative."
    )

# --- Retrieval Settings (New Expander) ---
with st.sidebar.expander("Retrieval Settings", expanded=True):
    st.markdown("Control how relevant documents are retrieved from the knowledge base (used by KnowledgeBase_Search tool).")

    st.session_state.top_k_retrieval = st.slider(
        "Number of Chunks (k)",
        min_value=1,
        max_value=10,
        value=4,
        step=1,
        help="Number of top similar document chunks to retrieve."
    )

    st.session_state.selected_search_type = st.selectbox(
        "Search Type",
        options=["Similarity", "MMR"], # MMR will act as Similarity for FAISS
        index=0,
        help="Similarity: Retrieves most similar chunks. MMR: (Note: For FAISS, this will behave like Similarity due to direct support limitations)."
    )

    if st.session_state.selected_search_type == "MMR":
        st.markdown("MMR Parameters (Note: For FAISS, these parameters are not directly used, search will be similarity-based):")
        mmr_col1, mmr_col2 = st.columns(2)
        with mmr_col1:
            st.session_state.fetch_k_mmr = st.number_input(
                "Fetch K (MMR)",
                min_value=st.session_state.top_k_retrieval,
                max_value=50,
                value=max(st.session_state.top_k_retrieval, 20),
                step=1,
                help="Ignored for FAISS; for compatibility with MMR selection."
            )
        with mmr_col2:
            st.session_state.lambda_mult_mmr = st.slider(
                "Lambda (MMR)",
                min_value=0.0,
                max_value=1.0,
                value=0.5,
                step=0.05,
                help="Ignored for FAISS; for compatibility with MMR selection."
            )
    else:
        if "fetch_k_mmr" not in st.session_state:
             st.session_state.fetch_k_mmr = 20
        if "lambda_mult_mmr" not in st.session_state:
            st.session_state.lambda_mult_mmr = 0.5


# --- Load from URL (using columns for input and button) ---
st.sidebar.markdown("---")
st.sidebar.subheader("Load from URL")
url_col, button_col = st.sidebar.columns([3, 1])

with url_col:
    # Use a dynamic key for the text input to allow easy clearing
    url_input_text = st.text_input(
        "Enter URL",
        label_visibility="collapsed",
        placeholder="Enter a URL to load (e.g., https://example.com)...",
        key=f"url_input_{st.session_state.url_input_key}", # Dynamic key based on counter
    )

with button_col:
    load_url_button_clicked = st.button("Load", key="load_url_button", use_container_width=True)


# --- Handle URL Button Click to Set submitted_url in Session State ---
if load_url_button_clicked:
    if url_input_text:
        st.session_state.submitted_url = url_input_text
        # Increment key to clear the visible input field on rerun
        st.session_state.url_input_key += 1
        st.session_state.uploaded_file_key += 1 # Clear file uploader too if new URL is loaded
        st.session_state.last_uploaded_filename = "" # Reset file status
        st.session_state.messages = [] # Clear chat on new content load
        st.rerun() # Force rerun to process submitted_url
    else:
        st.warning("Please enter a URL before clicking Load.")
        st.session_state.submitted_url = "" # Ensure it's empty if button clicked with no URL
        st.session_state.last_loaded_url = "" # Also clear last loaded on empty submission
        st.session_state.url_input_key += 1
        st.rerun() # Force rerun to update UI with warning


# --- Initial Load/Check for Existing DB ---
# This block always calls manage_vector_store on app startup/rerun
# The UI feedback (spinner/toast) for embedding model loading is now handled inside manage_vector_store
manage_vector_store(collection_name=st.session_state.current_collection_name)


# --- Handle Document Upload (triggered by file_uploader) ---
if uploaded_file and uploaded_file.name != st.session_state.get("last_uploaded_filename", ""):
    st.session_state.messages = []
    st.session_state.last_uploaded_filename = uploaded_file.name
    st.session_state.submitted_url = "" # Clear submitted URL on file upload
    st.session_state.last_loaded_url = "" # Clear last loaded URL on file upload
    st.session_state.url_input_key += 1 # Clear URL input field

    with st.spinner(f"Processing document '{uploaded_file.name}'..."):
        text_chunks = process_document_to_chunks(uploaded_file, st.session_state.chunk_size, st.session_state.chunk_overlap)

    if text_chunks:
        manage_vector_store(text_chunks=text_chunks, collection_name=clean_collection_name(os.path.splitext(uploaded_file.name)[0]))
        st.toast(f"Knowledge base ready for '{uploaded_file.name}'! You can now ask questions.", icon="🎉")
        st.rerun()
    else:
        st.session_state.vector_db = None
        st.session_state.current_content_source = None
        st.session_state.uploaded_file_key += 1 # Increment key to force re-render of file uploader if needed
        st.rerun()


# --- Handle URL Processing (triggered by submitted_url in session state) ---
# This block runs in a subsequent rerun after the button click sets submitted_url
if st.session_state.submitted_url and st.session_state.submitted_url != st.session_state.last_loaded_url:
    url_to_process = st.session_state.submitted_url
    
    text_chunks = process_url_to_chunks(url_to_process, st.session_state.chunk_size, st.session_state.chunk_overlap)

    if text_chunks:
        manage_vector_store(text_chunks=text_chunks, collection_name=get_url_collection_name(url_to_process))
        st.toast(f"Knowledge base ready for '{url_to_process}'! You can now ask questions.", icon="🎉")
        st.session_state.last_loaded_url = url_to_process # Update last loaded successfully
        st.session_state.submitted_url = "" # Clear submitted_url to prevent re-processing
        # No st.rerun() here, as manage_vector_store already triggers needed updates implicitly or the change in session state does.
        # However, if chat history needs clearing, we already do it on button click.
        # If any other part of the UI needs a hard refresh, a st.rerun() could be added, but let's avoid it unless necessary.
    else:
        st.session_state.vector_db = None
        st.session_state.current_content_source = None
        st.session_state.last_loaded_url = "" # Clear last loaded URL on failure
        st.session_state.submitted_url = "" # Clear submitted_url on failure
        st.rerun() # Force a rerun to update UI on failure


# --- Display Current Knowledge Base Status ---
if st.session_state.vector_db is None:
    st.toast("Please upload a document or load a URL to begin.", icon="⬆️")
    st.write("Current status: No knowledge base loaded.")
elif st.session_state.current_content_source:
    st.write(f"Knowledge base active for: **{st.session_state.current_content_source}** (Collection: `{st.session_state.current_collection_name}`)")
else:
    st.toast("No content loaded yet. Please upload a document or load a URL to begin.", icon="🤷‍♀️")
    st.write("Current status: Awaiting content.")


# --- Chat Interface ---
if not st.session_state.messages:
    initial_ai_response = "Hi there! I'm your RAG chatbot. I can answer questions about the document or URL you provide, and now also search the web for general knowledge!"
    st.session_state.messages = [{"role": "assistant", "content": initial_ai_response}]

for message in st.session_state.messages:
    with st.chat_message(message["role"]):
        st.write(message["content"])

if query := st.chat_input("Ask a question..."):
    st.session_state.messages.append({"role": "user", "content": query})
    with st.chat_message("user"):
        st.write(query)

    with st.spinner("Thinking... (The agent is deciding whether to use the knowledge base or web search...)"):
        response_text, source_docs = generate_answer_with_memory(
            query,
            st.session_state.messages[:-1], # Pass all but the current user query for history
            st.session_state.vector_db
        )

    st.session_state.messages.append({"role": "assistant", "content": response_text})
    with st.chat_message("assistant"):
        st.write(response_text)


# --- Sidebar: Clear Chat and Reset RAG ---
st.sidebar.markdown("---")
if st.sidebar.button("Clear Chat and Reset RAG Data"):
    st.session_state.messages = []
    st.session_state.vector_db = None
    st.session_state.current_content_source = None
    st.session_state.current_collection_name = DEFAULT_COLLECTION_NAME

    # Increment keys to clear input fields
    st.session_state.url_input_key += 1 # Increment to clear URL text input
    st.session_state.uploaded_file_key += 1 # Increment to clear file uploader
    st.session_state.last_uploaded_filename = ""
    st.session_state.last_loaded_url = ""
    st.session_state.submitted_url = "" # Clear submitted URL on reset

    if 'faiss_indexes' in st.session_state:
        st.session_state.faiss_indexes = {}
        st.toast("In-memory knowledge bases cleared!", icon="🗑️")

    st.rerun() # Force a full rerun to reflect all changes